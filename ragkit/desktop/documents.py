"""Folder scanning and document analysis utilities."""

from __future__ import annotations

from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import datetime, timezone
import fnmatch
import hashlib
import json
from pathlib import Path
import re
import unicodedata

from .models import (
    DocumentInfo,
    FolderEntry,
    FolderScanResult,
    FolderStats,
    FolderValidationResult,
    IngestionConfig,
    ScanFolderRequest,
    SUPPORTED_FILE_TYPES,
    FileTypeInfo,
    FolderNode,
)
from .analysis_progress import AnalysisProgress

import os

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import yaml
except ImportError:
    yaml = None

try:
    from langdetect import detect as langdetect_detect
except ImportError:
    langdetect_detect = None

try:
    import olefile
except ImportError:
    olefile = None

import mimetypes

@dataclass
class ParsedContent:
    text: str
    page_count: int | None
    title: str | None
    author: str | None
    creation_date: str | None
    encoding: str | None
    has_tables: bool = False
    has_images: bool = False
    has_code: bool = False
    parser_engine: str | None = None
    ocr_applied: bool = False

SUPPORTED_DISPLAY_NAMES = {
    "pdf": "PDF",
    "docx": "Word",
    "doc": "Word",
    "md": "Markdown",
    "txt": "Texte",
    "html": "HTML",
    "csv": "CSV",
    "rst": "ReStructuredText",
    "xml": "XML",
    "json": "JSON",
    "yaml": "YAML",
}

STOPWORDS = {
    "the", "a", "an", "and", "or", "but", "in", "on", "at", "to", "for", "of", "with", "by",
    "le", "la", "les", "un", "une", "des", "et", "ou", "mais", "dans", "sur", "à", "pour", "de", "avec", "par",
    "ce", "cette", "ces", "qui", "que", "quoi", "dont", "où", "is", "are", "was", "were", "be", "been",
}

def build_tree(path: Path, max_depth: int = 2, current_depth: int = 0) -> FolderNode:
    node = FolderNode(name=path.name, path=str(path), is_dir=True)
    
    if current_depth >= max_depth:
        return node

    try:
        # Sort directories first, then files
        entries = sorted(list(os.scandir(path)), key=lambda e: e.name.lower())
        file_count = 0
        size_bytes = 0
        
        for entry in entries:
            if entry.is_dir(follow_symlinks=False):
                if entry.name.startswith("."): 
                    continue
                    
                child = build_tree(Path(entry.path), max_depth, current_depth + 1)
                node.children.append(child)
                file_count += child.file_count 
                size_bytes += child.size_bytes
            elif entry.is_file(follow_symlinks=False):
                file_count += 1
                try:
                    size_bytes += entry.stat().st_size
                except OSError:
                    pass
        
        node.file_count = file_count
        node.size_bytes = size_bytes

    except PermissionError:
        pass
    except OSError:
        pass

    return node

def validate_folder(folder_path: str, recursive: bool = True) -> FolderValidationResult:
    root = Path(folder_path).expanduser()
    if not root.exists():
        return FolderValidationResult(
            valid=False,
            error="Folder does not exist.",
            error_code="FOLDER_NOT_FOUND",
            stats=FolderStats(files=0, size_mb=0.0, extensions=[], extension_counts={}),
            subdirectories=[],
        )
    if not root.is_dir():
        return FolderValidationResult(
            valid=False,
            error="Path is not a directory.",
            error_code="NOT_A_DIRECTORY",
            stats=FolderStats(files=0, size_mb=0.0, extensions=[], extension_counts={}),
            subdirectories=[],
        )

    # For stats, we use the existing _iter_files logic but we need to respect recursive flag passed from wizard
    # The existing validate_folder didn't take recursive arg, but wizard needs it.
    # We'll default to True to maintain compatibility or update signature.
    # The spec for wizard says validation returns stats.
    
    files = list(_iter_files(root, recursive=recursive, excluded_dirs=[], exclusion_patterns=[], max_file_size_mb=None))
    stats = _build_folder_stats(root, files)
    subdirectories = _build_subdirectory_stats(root, files)

    # Build tree for UI (limited depth to avoid huge payload)
    tree_root = build_tree(root, max_depth=3)

    return FolderValidationResult(
        valid=True, 
        stats=stats, 
        subdirectories=subdirectories,
        tree=tree_root
    )


def scan_folder(payload: ScanFolderRequest) -> FolderScanResult:
    root = Path(payload.folder_path).expanduser()
    if not root.exists() or not root.is_dir():
        return FolderScanResult(
            supported_types=[],
            unsupported_types=[],
            total_files=0,
            total_size_mb=0.0,
        )

    files = list(
        _iter_files(
            root,
            recursive=payload.recursive,
            excluded_dirs=payload.excluded_dirs,
            exclusion_patterns=payload.exclusion_patterns,
            max_file_size_mb=payload.max_file_size_mb,
        )
    )

    counters: dict[str, int] = defaultdict(int)
    sizes: dict[str, int] = defaultdict(int)
    total_size = 0

    for path in files:
        extension = _normalize_extension(path.suffix)
        counters[extension] += 1
        size = path.stat().st_size
        sizes[extension] += size
        total_size += size

    supported_types: list[FileTypeInfo] = []
    unsupported_types: list[FileTypeInfo] = []

    for extension, count in sorted(counters.items(), key=lambda item: (-item[1], item[0])):
        info = FileTypeInfo(
            extension=extension,
            display_name=SUPPORTED_DISPLAY_NAMES.get(extension, extension.upper()),
            count=count,
            size_mb=round(sizes[extension] / (1024 * 1024), 2),
            supported=extension in SUPPORTED_FILE_TYPES,
        )
        if info.supported:
            supported_types.append(info)
        else:
            unsupported_types.append(info)

    return FolderScanResult(
        supported_types=supported_types,
        unsupported_types=unsupported_types,
        total_files=sum(counters.values()),
        total_size_mb=round(total_size / (1024 * 1024), 2),
    )


def analyze_documents(config: IngestionConfig) -> tuple[list[DocumentInfo], list[str]]:
    source = config.source
    root = Path(source.path).expanduser()
    if not root.exists() or not root.is_dir():
        return [], ["Source folder does not exist."]

    selected_types = {_normalize_extension(file_type) for file_type in source.file_types}
    files = []
    for path in _iter_files(
        root,
        recursive=source.recursive,
        excluded_dirs=source.excluded_dirs,
        exclusion_patterns=source.exclusion_patterns,
        max_file_size_mb=source.max_file_size_mb,
    ):
        if _normalize_extension(path.suffix) in selected_types:
            files.append(path)

    parsed_documents: list[DocumentInfo] = []
    errors: list[str] = []
    seen_hashes: set[str] = set()
    seen_token_sets: list[set[str]] = []

    progress = AnalysisProgress.get_instance()
    progress.start(total=len(files))

    for file_path in files:
        try:
            progress.update(current_file=file_path.name)
            parsed = _extract_content(file_path)
            preprocessed = _preprocess_text(parsed.text, config)
            if _is_duplicate(preprocessed, config, seen_hashes, seen_token_sets):
                continue

            relative_path = file_path.relative_to(root).as_posix()
            file_type = _normalize_extension(file_path.suffix)
            detected_language = _detect_language(preprocessed) if config.preprocessing.language_detection else None
            title, description = _derive_title_description(preprocessed, file_path.stem, parsed.title)
            keywords = _extract_keywords(preprocessed)

            parsed_documents.append(
                DocumentInfo(
                    id=hashlib.sha1(relative_path.encode("utf-8")).hexdigest(),
                    filename=file_path.name,
                    file_path=relative_path,
                    file_type=file_type,
                    file_size_bytes=file_path.stat().st_size,
                    page_count=parsed.page_count,
                    language=detected_language,
                    last_modified=datetime.fromtimestamp(
                        file_path.stat().st_mtime, tz=timezone.utc
                    ).isoformat(),
                    encoding=parsed.encoding,
                    word_count=len(re.findall(r"\w+", preprocessed)),
                    title=title,
                    author=parsed.author,
                    description=description,
                    keywords=keywords,
                    creation_date=parsed.creation_date,
                    mime_type=mimetypes.guess_type(file_path)[0],
                    ingested_at=datetime.now(timezone.utc).isoformat(),
                    char_count=len(parsed.text),
                    has_tables=parsed.has_tables,
                    has_images=parsed.has_images,
                    has_code=parsed.has_code,
                    parser_engine=parsed.parser_engine,
                    ocr_applied=parsed.ocr_applied,
                    tags=keywords,
                    category=None,
                    text_preview=parsed.text[:500] if parsed.text else None,
                )
            )
        except Exception as exc:  # pragma: no cover - defensive at runtime
            progress.update(current_file=file_path.name, error=True)
            errors.append(f"{file_path.name}: {exc}")
    
    progress.finish()
    return parsed_documents, errors


def _build_folder_stats(root: Path, files: list[Path]) -> FolderStats:
    extension_counts: dict[str, int] = defaultdict(int)
    total_size = 0
    for path in files:
        extension = _normalize_extension(path.suffix)
        extension_counts[extension] += 1
        total_size += path.stat().st_size
    return FolderStats(
        files=len(files),
        size_mb=round(total_size / (1024 * 1024), 2),
        extensions=sorted(extension_counts.keys()),
        extension_counts=dict(sorted(extension_counts.items())),
    )


def _build_subdirectory_stats(root: Path, files: list[Path]) -> list[FolderEntry]:
    counters: dict[str, int] = defaultdict(int)
    for file_path in files:
        relative = file_path.relative_to(root)
        if len(relative.parts) <= 1:
            continue
        directory = "/".join(relative.parts[:-1])
        counters[directory] += 1
    return [FolderEntry(path=path, files=count) for path, count in sorted(counters.items())]


def _iter_files(
    root: Path,
    recursive: bool,
    excluded_dirs: list[str],
    exclusion_patterns: list[str],
    max_file_size_mb: int | None,
):
    excluded: set[str] = set()
    for value in excluded_dirs:
        if not value.strip():
            continue
        p = Path(value)
        if p.is_absolute():
            try:
                excluded.add(p.relative_to(root).as_posix().lower())
            except ValueError:
                excluded.add(_normalize_relative_path(value))
        else:
            excluded.add(_normalize_relative_path(value))
    max_size_bytes = max_file_size_mb * 1024 * 1024 if max_file_size_mb else None
    iterator = root.rglob("*") if recursive else root.glob("*")

    for path in iterator:
        if not path.is_file():
            continue
        relative = path.relative_to(root)
        relative_str = relative.as_posix()

        if path.name.startswith("~$"):
            continue
        if _is_excluded(relative_str, excluded):
            continue
        if _matches_patterns(path.name, relative_str, exclusion_patterns):
            continue
        if max_size_bytes is not None and path.stat().st_size > max_size_bytes:
            continue
        yield path


def _normalize_relative_path(value: str) -> str:
    return value.replace("\\", "/").strip("/").lower()


def _is_excluded(relative_path: str, excluded_dirs: set[str]) -> bool:
    lowered = relative_path.lower()
    for excluded in excluded_dirs:
        if lowered == excluded or lowered.startswith(f"{excluded}/"):
            return True
    return False


def _matches_patterns(filename: str, relative_path: str, patterns: list[str]) -> bool:
    for pattern in patterns:
        if not pattern.strip():
            continue
        if fnmatch.fnmatch(filename, pattern) or fnmatch.fnmatch(relative_path, pattern):
            return True
    return False


def _normalize_extension(extension: str) -> str:
    normalized = extension.lower().lstrip(".")
    return "yaml" if normalized == "yml" else normalized


def _preprocess_text(text: str, config: IngestionConfig) -> str:
    processed = text
    preprocessing = config.preprocessing

    if preprocessing.normalize_unicode:
        processed = unicodedata.normalize("NFC", processed)
    if preprocessing.remove_urls:
        processed = re.sub(r"https?://\S+", " ", processed)
    if preprocessing.remove_punctuation:
        processed = re.sub(r"[^\w\s]", " ", processed)
    if preprocessing.lowercase:
        processed = processed.lower()

    processed = re.sub(r"\s+", " ", processed).strip()
    return processed


def _is_duplicate(
    text: str,
    config: IngestionConfig,
    seen_hashes: set[str],
    seen_token_sets: list[set[str]],
) -> bool:
    strategy = config.preprocessing.deduplication_strategy.value
    threshold = config.preprocessing.deduplication_threshold
    if not text:
        return True
    if strategy == "none":
        return False

    text_hash = hashlib.sha256(text.encode("utf-8")).hexdigest()
    if strategy == "exact":
        if text_hash in seen_hashes:
            return True
        seen_hashes.add(text_hash)
        return False

    current_tokens = set(re.findall(r"\w+", text.lower()))
    if not current_tokens:
        return False
    for existing in seen_token_sets:
        score = _jaccard_similarity(existing, current_tokens)
        if score >= threshold:
            return True
    seen_token_sets.append(current_tokens)
    return False


def _jaccard_similarity(left: set[str], right: set[str]) -> float:
    union = left | right
    if not union:
        return 0.0
    return len(left & right) / len(union)


def _detect_language(text: str) -> str | None:
    if not text or langdetect_detect is None:
        return None
    sample = text[:5000]
    try:
        return langdetect_detect(sample)
    except Exception:  # pragma: no cover - runtime dependency behavior
        return None


def _derive_title_description(text: str, fallback_title: str, extracted_title: str | None) -> tuple[str, str | None]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    title = extracted_title or (lines[0] if lines else fallback_title)
    if len(title) > 160:
        title = title[:157] + "..."
    if not lines:
        return title, None
    description = " ".join(lines[:2])
    if len(description) > 300:
        description = description[:297] + "..."
    return title, description


def _extract_keywords(text: str, max_keywords: int = 8) -> list[str]:
    words = re.findall(r"[A-Za-zÀ-ÿ]{4,}", text.lower())
    filtered = [word for word in words if word not in STOPWORDS]
    counts = Counter(filtered)
    return [word for word, _count in counts.most_common(max_keywords)]


def _extract_content(path: Path) -> ParsedContent:
    # 1. Check for empty file
    if path.stat().st_size == 0:
        return ParsedContent(text="", page_count=None, title=None, author=None, creation_date=None, encoding=None)

    # 2. Check for OLE2 legacy Word format
    try:
        with open(path, "rb") as f:
            header = f.read(8)
            if header == b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1":
                return _extract_doc_legacy(path)
    except Exception:
        pass

    extension = _normalize_extension(path.suffix)
    if extension == "pdf":
        return _extract_pdf(path)
    if extension in {"docx", "doc"}:
        return _extract_docx(path)
    if extension == "json":
        return _extract_json(path)
    if extension in {"yaml", "yml"}:
        return _extract_yaml(path)
    if extension == "html":
        return _extract_html(path)
    return _extract_text(path)


def _extract_doc_legacy(path: Path) -> ParsedContent:
    if olefile is None:
        # Fallback if olefile is missing (though it should be installed)
        return ParsedContent(text="", page_count=None, title=None, author=None, creation_date=None, encoding=None, parser_engine="olefile-missing")
    
    text = ""
    title = None
    author = None
    creation_date = None
    
    try:
        with olefile.OleFileIO(path) as ole:
            # Try to extract metadata from properties
            meta = ole.get_metadata()
            if meta:
                title = meta.title
                author = meta.author
                if meta.create_time:
                    creation_date = meta.create_time.isoformat()

            # Extract text from WordDocument stream
            if ole.exists("WordDocument"):
                with ole.openstream("WordDocument") as stream:
                    data = stream.read()
                    # Basic string extraction: sequences of 4+ printable chars
                    # 1. Try to find UTF-16LE strings (common in Word binary)
                    # Pattern: [char][\x00] repeated 4+ times
                    utf16_strings = re.findall(b"(?:[\x20-\x7E][\x00]){4,}", data)
                    
                    # 2. Try to find ASCII/CP1252 strings
                    ascii_strings = re.findall(b"[\x20-\x7E]{4,}", data)
                    
                    # Decode and join
                    decoded = []
                    for s in utf16_strings:
                        try:
                            decoded.append(s.decode("utf-16le"))
                        except:
                            pass
                    for s in ascii_strings:
                        try:
                            decoded.append(s.decode("cp1252"))
                        except:
                            pass
                            
                    text = "\n".join(decoded)
                    
                    # Clean up common binary artifacts
                    # "bjbj" is a common header in Word OLE streams
                    text = re.sub(r"bjbj\w+", "", text)
                    text = re.sub(r"\x00+", "", text)
                    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
                    
    except Exception:
        pass

    return ParsedContent(
        text=text,
        page_count=None,
        title=_validate_extracted_title(title, path.stem),
        author=author,
        creation_date=creation_date,
        encoding="ole2-binary",
        parser_engine="olefile",
        has_tables=False, 
        has_images=False,
    )

def _validate_extracted_title(title: str | None, filename_stem: str) -> str | None:
    if not title:
        return None
        
    s = title.strip()
    if not s:
        return None
        
    # Check for binary artifacts / garbage
    if s.lower().startswith("bjbj"):
        return None
        
    # Check for high ratio of non-printable/special chars
    # If < 50% of chars are alphanumeric/space, it's likely garbage
    alnum_count = sum(c.isalnum() or c.isspace() for c in s)
    if len(s) > 0 and (alnum_count / len(s)) < 0.5:
        return None
        
    # Check if title is just the filename (redundant)
    # or if it looks like a filename (extension)
    if s.lower() == filename_stem.lower():
        return None
        
    return s


def _extract_pdf(path: Path) -> ParsedContent:
    if PdfReader is None:  # pragma: no cover - optional dependency branch
        return ParsedContent(text="", page_count=None, title=None, author=None, creation_date=None, encoding=None)

    reader = PdfReader(str(path))
    page_count = len(reader.pages)
    text_chunks = []
    has_images = False
    for page in reader.pages:
        text_chunks.append(page.extract_text() or "")
        if not has_images and '/XObject' in page: # simplified check
             # Deep check requires iterating resources
             pass 
        # pypdf images access is expensive, maybe skip for speed or do better check if reliable
        try:
            if len(page.images) > 0:
                has_images = True
        except:
            pass

    metadata = reader.metadata or {}

    return ParsedContent(
        text="\n".join(text_chunks),
        page_count=page_count,
        title=_clean_metadata_value(getattr(metadata, "title", None)),
        author=_clean_metadata_value(getattr(metadata, "author", None)),
        creation_date=_clean_metadata_value(getattr(metadata, "creation_date", None)),
        encoding=None,
        has_images=has_images,
        parser_engine="pypdf",
    )


def _extract_docx(path: Path) -> ParsedContent:
    if DocxDocument is None:  # pragma: no cover - optional dependency branch
        return ParsedContent(text="", page_count=None, title=None, author=None, creation_date=None, encoding=None)

    document = DocxDocument(str(path))
    text = "\n".join([paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()])
    words = re.findall(r"\w+", text)
    page_count = max(1, round(len(words) / 500)) if words else 1

    core = document.core_properties
    creation_date = core.created.isoformat() if core.created else None

    has_tables = len(document.tables) > 0
    has_images = len(document.inline_shapes) > 0 # basic check for inline images

    return ParsedContent(
        text=text,
        page_count=page_count,
        title=core.title or None,
        author=core.author or None,
        creation_date=creation_date,
        encoding="utf-8",
        has_tables=has_tables,
        has_images=has_images,
        parser_engine="python-docx",
    )


def _extract_json(path: Path) -> ParsedContent:
    text, encoding = _read_text_file(path)
    try:
        parsed = json.loads(text)
        text = json.dumps(parsed, ensure_ascii=False, indent=2)
    except Exception:
        pass
    return ParsedContent(text=text, page_count=None, title=None, author=None, creation_date=None, encoding=encoding, parser_engine="json", has_code=True)


def _extract_yaml(path: Path) -> ParsedContent:
    text, encoding = _read_text_file(path)
    if yaml is not None:
        try:
            parsed = yaml.safe_load(text)
            text = json.dumps(parsed, ensure_ascii=False, indent=2) if parsed is not None else ""
        except Exception:
            pass
    return ParsedContent(text=text, page_count=None, title=None, author=None, creation_date=None, encoding=encoding, parser_engine="yaml", has_code=True)


def _extract_html(path: Path) -> ParsedContent:
    text, encoding = _read_text_file(path)
    stripped = re.sub(r"<script[\s\S]*?</script>", " ", text, flags=re.IGNORECASE)
    stripped = re.sub(r"<style[\s\S]*?</style>", " ", stripped, flags=re.IGNORECASE)
    stripped = re.sub(r"<[^>]+>", " ", stripped)
    return ParsedContent(
        text=re.sub(r"\s+", " ", stripped),
        page_count=None,
        title=None,
        author=None,
        creation_date=None,
        encoding=encoding,
        parser_engine="html",
        has_code=True,
    )


def _extract_text(path: Path) -> ParsedContent:
    text, encoding = _read_text_file(path)
    has_code = path.suffix.lower() in {'.py', '.js', '.ts', '.rs', '.java', '.cpp', '.c', '.h', '.css', '.sh', '.bat'}
    return ParsedContent(text=text, page_count=None, title=None, author=None, creation_date=None, encoding=encoding, parser_engine="text", has_code=has_code)


def _read_text_file(path: Path) -> tuple[str, str | None]:
    raw = path.read_bytes()
    for encoding in ("utf-8", "utf-16", "cp1252", "latin-1"):
        try:
            return raw.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore"), None


def _clean_metadata_value(value: object) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    if text.startswith("D:") and len(text) >= 10:
        # Common PDF metadata format: D:YYYYMMDDHHmmSS...
        digits = re.sub(r"[^\d]", "", text[2:])
        if len(digits) >= 8:
            year = int(digits[0:4])
            month = int(digits[4:6])
            day = int(digits[6:8])
            try:
                return datetime(year, month, day, tzinfo=timezone.utc).isoformat()
            except ValueError:
                return text
    return text
